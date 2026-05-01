"""Word proposal builder."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config import OUTPUTS_DIR
from utils.deployment_utils import sanitize_mission_name


DISCLAIMER_TEXT = (
    "This document is an AI-assisted Phase 0 concept draft. All technical values, cost estimates, "
    "procurement options, and requirements must be reviewed and validated by a qualified spacecraft systems "
    "engineer before use in design, procurement, proposal submission, or mission decision-making."
)
MVP_SCOPE_1 = (
    "This MVP report is generated for a single-satellite mission concept. It does not analyze constellation "
    "deployment, multi-satellite phasing, revisit optimization, inter-satellite links, or multi-plane architectures."
)
MVP_SCOPE_2 = (
    "This MVP report assumes an Earth Observation payload concept. Payload-specific analysis is limited to high-level "
    "swath, resolution, mass, and power assumptions where provided."
)


class ProposalBuilder:
    """Build Phase 0 proposal document from coordinator outputs."""

    VERSION = "MVP Draft"

    def build_proposal(self, mission_context: dict, outputs: dict) -> Path:
        OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
        doc = Document()
        self._configure_document(doc)
        self._add_header_footer(doc, mission_context["mission_name"])

        mission = outputs.get("mission", {})
        mass = outputs.get("mass", {})
        power = outputs.get("power", {})
        cost = outputs.get("cost", {})
        procurement = outputs.get("procurement", {})

        doc.add_heading(mission_context["mission_name"], level=0)
        doc.add_paragraph("Phase 0 Mission Proposal")
        doc.add_paragraph(f"Date: {date.today().isoformat()}")
        doc.add_paragraph(f"Version: {self.VERSION}")
        doc.add_page_break()

        doc.add_heading("MVP Scope and Limitations", level=1)
        doc.add_paragraph(f"1. {MVP_SCOPE_1}")
        doc.add_paragraph(f"2. {MVP_SCOPE_2}")

        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(mission.get("conops_summary", "Executive summary unavailable."))

        doc.add_heading("Mission Overview", level=1)
        for obj in mission.get("objectives", ["No objectives available."]):
            doc.add_paragraph(f"- {obj}")

        doc.add_heading("Mission Requirements", level=1)
        req_table = doc.add_table(rows=1, cols=3)
        self._style_table(req_table, ["ID", "Requirement", "Rationale"])
        for req in mission.get("requirements", []):
            row = req_table.add_row().cells
            row[0].text = str(req.get("id", ""))
            row[1].text = str(req.get("text", ""))
            row[2].text = str(req.get("rationale", ""))

        doc.add_heading("Mass Budget", level=1)
        self._add_status_line(doc, "Mass", mass.get("status", "UNKNOWN"))
        mass_table = doc.add_table(rows=1, cols=4)
        self._style_table(mass_table, ["Subsystem", "Mass (kg)", "Margin %", "Notes"])
        for sub in mass.get("subsystems", []):
            row = mass_table.add_row().cells
            row[0].text = str(sub.get("name", ""))
            row[1].text = str(sub.get("mass_kg", ""))
            row[2].text = str(sub.get("margin_pct", ""))
            row[3].text = str(sub.get("notes", ""))
        doc.add_paragraph(f"Total dry mass (kg): {mass.get('total_dry_mass_kg', 'N/A')}")
        doc.add_paragraph(f"Total wet mass (kg): {mass.get('total_wet_mass_kg', 'N/A')}")

        doc.add_heading("Power Budget", level=1)
        if power.get("_fallback"):
            doc.add_paragraph(
                f"WARNING: Power section is fallback output due to agent failure: {power.get('_error', 'Unknown error')}"
            )
        self._add_status_line(doc, "Power", power.get("status", "UNKNOWN"))
        for mode, mode_data in power.get("modes", {}).items():
            doc.add_paragraph(f"Mode: {mode} - Total W: {mode_data.get('total_W', 'N/A')}")
        doc.add_paragraph(f"Solar array area (m^2): {power.get('solar_array_area_m2', 'N/A')}")
        doc.add_paragraph(f"Battery capacity (Wh): {power.get('battery_capacity_Wh', 'N/A')}")
        doc.add_paragraph(f"Eclipse duration (min): {power.get('eclipse_duration_min', 'N/A')}")

        doc.add_heading("Cost Estimate", level=1)
        if cost.get("_fallback"):
            doc.add_paragraph(
                f"WARNING: Cost section is fallback output due to agent failure: {cost.get('_error', 'Unknown error')}"
            )
        cost_table = doc.add_table(rows=1, cols=3)
        self._style_table(cost_table, ["Category", "Cost (kEUR)", "Basis"])
        for item in cost.get("cost_breakdown", []):
            row = cost_table.add_row().cells
            row[0].text = str(item.get("category", ""))
            row[1].text = str(item.get("cost_kEUR", ""))
            row[2].text = str(item.get("basis", ""))
        doc.add_paragraph(f"Total cost (kEUR): {cost.get('total_cost_kEUR', 'N/A')}")

        doc.add_heading("Component Alternatives", level=1)
        if procurement.get("_fallback"):
            doc.add_paragraph(
                "WARNING: Procurement section is fallback output due to upstream/agent failure."
            )
        for search in procurement.get("component_searches", []):
            doc.add_paragraph(f"Category: {search.get('category', '')}")
            comp_table = doc.add_table(rows=1, cols=6)
            self._style_table(comp_table, ["Rank", "Product", "Vendor", "URL", "Price kEUR", "Lead weeks"])
            for alt in search.get("alternatives", []):
                row = comp_table.add_row().cells
                row[0].text = str(alt.get("rank", ""))
                row[1].text = str(alt.get("product_name", ""))
                row[2].text = str(alt.get("vendor", ""))
                row[3].text = str(alt.get("source_url", ""))
                row[4].text = str(alt.get("unit_price_kEUR", ""))
                row[5].text = str(alt.get("lead_time_weeks", ""))
            for warn in search.get("warnings", []):
                doc.add_paragraph(f"Warning: {warn}")

        doc.add_heading("Assumptions and Open Items", level=1)
        doc.add_paragraph("Mission assumptions:")
        for a in mission.get("assumptions", []):
            doc.add_paragraph(f"- {a}")
        doc.add_paragraph("Open questions:")
        for q in mission.get("open_questions", []):
            doc.add_paragraph(f"- {q}")

        doc.add_heading("Disclaimer", level=1)
        doc.add_paragraph(DISCLAIMER_TEXT)

        safe_name = sanitize_mission_name(mission_context["mission_name"])
        path = OUTPUTS_DIR / f"Phase_0_Proposal_{safe_name}.docx"
        doc.save(path)
        return path

    def _configure_document(self, doc: Document) -> None:
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

    def _add_header_footer(self, doc: Document, mission_name: str) -> None:
        section = doc.sections[0]
        section.header.paragraphs[0].text = f"{mission_name} - {self.VERSION}"
        section.footer.paragraphs[0].text = "Page number placeholder"

    def _style_table(self, table, headers: list[str]) -> None:
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        table.style = "Table Grid"
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            shade = OxmlElement("w:shd")
            shade.set(qn("w:fill"), "D9D9D9")
            tc_pr.append(shade)

    def _add_status_line(self, doc: Document, label: str, status: str) -> None:
        p = doc.add_paragraph(f"{label} status: {status}")
        if not p.runs:
            return
        run = p.runs[0]
        if status == "GREEN":
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
        elif status == "YELLOW":
            run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
        elif status == "RED":
            run.font.color.rgb = RGBColor(0x99, 0x00, 0x00)
